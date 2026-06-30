import io
import re

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using pypdf."""
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return clean_text(text)
    except Exception as e:
        raise ValueError(f"Failed to extract PDF text: {str(e)}")

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return clean_text("\n".join(paragraphs))
    except Exception as e:
        raise ValueError(f"Failed to extract DOCX text: {str(e)}")

def clean_text(text: str) -> str:
    """Remove excessive whitespace while preserving structure."""
    # Replace multiple newlines with double newline
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Replace multiple spaces with single space
    text = re.sub(r' {2,}', ' ', text)
    return text.strip()
